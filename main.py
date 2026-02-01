# main.py (FINAL - ready to replace)
# Includes:
# - prompts registry import (backend/prompts/registry.py)
# - PDF upload size limit: 50MB (server-side hard limit)
# - Text chunking for >80k chars (chunk + merge)
# - Consistent fileUrl rules: only upload analyses have fileUrl; text analyses return None
# - prompt_version persisted + migrated
# - history returns promptVersion

from fastapi import FastAPI, HTTPException, Header, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from fastapi.staticfiles import StaticFiles
import uuid
from pydantic import BaseModel
from typing import Literal, List, Optional, Any, Dict, Tuple
import json, os, hashlib
import urllib.request
import urllib.parse
from datetime import datetime, timedelta
import re
import fitz  # pymupdf
from docx import Document
import io
from openai import OpenAI

from sqlalchemy import create_engine, Column, Integer, String, Text, DateTime, LargeBinary, delete
from sqlalchemy.orm import sessionmaker, declarative_base

from passlib.context import CryptContext
import jwt

from prompts.registry import PROMPT_VERSION, validate_contract_type, build_system_prompt

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(APP_DIR, "contract_ai.db")
DATABASE_URL = f"sqlite:///{DB_PATH}"

JWT_SECRET = os.getenv("JWT_SECRET", "change-me-now")
JWT_ALG = "HS256"
JWT_EXPIRE_HOURS = 24
WECHAT_APPID = os.getenv("WECHAT_APPID", "")
WECHAT_SECRET = os.getenv("WECHAT_SECRET", "")

# OpenAI
client = OpenAI()

pwd_context = CryptContext(schemes=["pbkdf2_sha256"], deprecated="auto")

engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()

# ---- Limits ----
MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50MB
MAX_TEXT_CHARS = 80_000
CHUNK_CHARS = 20_000  # per chunk
MAX_CHUNKS = 12        # safeguard (12 * 20k = 240k chars)


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True)
    phone = Column(String(32), unique=True, index=True, nullable=False)
    password_hash = Column(String(255), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    display_name = Column(String(64), nullable=True)
    avatar_url = Column(Text, nullable=True)
    credits = Column(Integer, nullable=False, default=0)


class Analysis(Base):
    __tablename__ = "analyses"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, index=True, nullable=False)

    contract_type = Column(String(100), nullable=False)
    identity = Column(String(1), nullable=False)

    prompt_version = Column(String(64), nullable=False, default=PROMPT_VERSION)

    content_hash = Column(String(64), index=True, nullable=False)

    original_content = Column(Text, nullable=False)
    result_json = Column(Text, nullable=False)

    file_name = Column(String(255), nullable=True)
    file_mime = Column(String(100), nullable=True)
    file_bytes = Column(LargeBinary, nullable=True)
    display_name = Column(String(255), nullable=True)

    created_at = Column(DateTime, default=datetime.utcnow)

class UploadBatch(Base):
    __tablename__ = "upload_batches"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, index=True, nullable=False)
    batch_id = Column(String(64), index=True, nullable=False)
    contract_type = Column(String(100), nullable=False)
    identity = Column(String(1), nullable=False)
    total = Column(Integer, nullable=False, default=1)
    received = Column(Integer, nullable=False, default=0)
    created_at = Column(DateTime, default=datetime.utcnow)


class UploadBatchFile(Base):
    __tablename__ = "upload_batch_files"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, index=True, nullable=False)
    batch_id = Column(String(64), index=True, nullable=False)
    idx = Column(Integer, nullable=False)  # 1..N
    file_name = Column(String(255), nullable=True)
    file_mime = Column(String(100), nullable=True)
    file_bytes = Column(LargeBinary, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)

class BatchFinalizeRequest(BaseModel):
    batch_id: str

# =========================
# Share (NEW)
# =========================
class AnalysisShare(Base):
    __tablename__ = "analysis_shares"

    id = Column(Integer, primary_key=True)

    # share_id: 给外部访问用，必须不可预测、唯一
    share_id = Column(String(64), unique=True, index=True, nullable=False)

    # 关联来源（内部追溯用，前端永远不需要知道）
    analysis_id = Column(Integer, index=True, nullable=False)
    user_id = Column(Integer, index=True, nullable=False)

    # Share 页需要展示的最小字段集（脱敏摘要）
    score = Column(Integer, nullable=False)
    score_title = Column(String(32), nullable=False)
    risk_summary = Column(Text, nullable=False)
    
    contract_name = Column(String(255), nullable=False)
    
    # 30 天有效期
    expires_at = Column(DateTime, index=True, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)


def _sqlite_add_column_if_missing(table: str, col: str, ddl_type: str):
    with engine.connect() as conn:
        rows = conn.exec_driver_sql(f"PRAGMA table_info({table});").fetchall()
        existing = {r[1] for r in rows}
        if col not in existing:
            conn.exec_driver_sql(f"ALTER TABLE {table} ADD COLUMN {col} {ddl_type};")


def _migrate_sqlite_schema():
    Base.metadata.create_all(bind=engine)
    _sqlite_add_column_if_missing("analyses", "file_name", "TEXT")
    _sqlite_add_column_if_missing("analyses", "file_mime", "TEXT")
    _sqlite_add_column_if_missing("analyses", "file_bytes", "BLOB")
    _sqlite_add_column_if_missing("analyses", "prompt_version", "TEXT")
    _sqlite_add_column_if_missing("analyses", "display_name", "TEXT")
    _sqlite_add_column_if_missing("upload_batches", "created_at", "DATETIME")
    _sqlite_add_column_if_missing("upload_batch_files", "created_at", "DATETIME")
    _sqlite_add_column_if_missing("users", "phone", "TEXT")
    _sqlite_add_column_if_missing("users", "display_name", "TEXT")
    _sqlite_add_column_if_missing("users", "avatar_url", "TEXT")
    _sqlite_add_column_if_missing("users", "credits", "INTEGER")    
    _sqlite_add_column_if_missing("analysis_shares", "contract_name", "TEXT")
    _sqlite_add_column_if_missing("analysis_shares", "created_at", "DATETIME")
    with engine.connect() as conn:
        conn.exec_driver_sql("CREATE UNIQUE INDEX IF NOT EXISTS ix_users_phone ON users (phone);")
        conn.exec_driver_sql("CREATE UNIQUE INDEX IF NOT EXISTS ix_analysis_shares_share_id ON analysis_shares (share_id);")
        conn.exec_driver_sql("CREATE INDEX IF NOT EXISTS ix_analysis_shares_expires_at ON analysis_shares (expires_at);")


_migrate_sqlite_schema()


app = FastAPI()

# ===== Static uploads (avatars) =====
UPLOAD_DIR = os.path.join(APP_DIR, "uploads")
AVATAR_DIR = os.path.join(UPLOAD_DIR, "avatars")
os.makedirs(AVATAR_DIR, exist_ok=True)

# 访问路径：/uploads/avatars/<file>
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

# CORS: allow localhost dev. Add prod front-end origin later if needed.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

 # [Dev (Vite)
        #"http://localhost:3000",
        #"http://127.0.0.1:3000",
        #"http://localhost:3001",
        #"http://127.0.0.1:3001",
        #"http://100.69.35.114:3000",
        #"http://100.69.35.114:3001",
        #"http://106.15.93.118:8000",
        #"http://106.15.93.118",
        #"https://safecontract.cn",
	#"http://safecontract.cn",
    #],

class RegisterRequest(BaseModel):
    phone: str
    password: str # 6位数字


class LoginRequest(BaseModel):
    phone: str
    password: str


class LoginResponse(BaseModel):
    access_token: str
    token_type: Literal["bearer"] = "bearer"

class WechatLoginRequest(BaseModel):
    code: str 

class RiskClause(BaseModel):
    section: str
    title: str
    originalText: str
    explanation: str
    suggestion: str
    level: Literal["HIGH", "MEDIUM", "LOW"]


class AnalyzeTextRequest(BaseModel):
    type: str = "general"
    identity: Literal["A", "B"]
    content: str


class AnalysisResult(BaseModel):
    id: str
    name: str
    date: str
    score: float
    riskSummary: str
    clauses: List[RiskClause]
    originalContent: str
    status: Literal["completed"]
    type: str
    identity: Literal["A", "B"]
    promptVersion: str
    imagePreview: Optional[str] = None
    fileUrl: Optional[str] = None
    fileName: Optional[str] = None
    meta: Optional[Dict[str, Any]] = None

OUTPUT_TEMPLATE = {"score": 0, "riskSummary": "", "originalContent": "", "clauses": []}

_CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")

# =========================
# Share Schemas (NEW)
# =========================
class CreateShareRequest(BaseModel):
    analysis_id: str  # 前端传 string，后端转 int
    contractName: Optional[str] = None

class CreateShareResponse(BaseModel):
    shareId: str
    expiresAt: str

class SharePublicResponse(BaseModel):
    contractName: str
    score: int
    scoreTitle: str
    riskSummary: str
    expiresAt: str

def _validate_phone_password(phone: str, password: str):
    p = (phone or "").strip()
    pw = (password or "").strip()

    # 手机号：先做最宽松的 11 位数字（中国）
    if not (len(p) == 11 and p.isdigit()):
        raise HTTPException(status_code=400, detail="Invalid phone (must be 11 digits)")

    # 密码：6位数字
    if not (len(pw) == 6 and pw.isdigit()):
        raise HTTPException(status_code=400, detail="Password must be 6 digits")

    return p, pw

CONTRACT_TYPE_CN = {
    "general": "通用合同",
    "marriage": "婚姻财产",
    "house_sale": "房屋买卖",
    "vehicle_sale": "车辆买卖",
    "lease": "租赁相关",
    "employment": "劳动合同",
    "nda": "保密协议",
    "service": "采购服务",
}

def _analysis_display_name(db, row: Analysis) -> str:
    """
    用 created_at 在同用户同类型里计算稳定序号，尽量还原：通用合同-01 / 通用合同-02 ...
    """
    base = CONTRACT_TYPE_CN.get(row.contract_type, "合同")
    if not row.created_at:
        return f"{base}-{int(row.id):02d}"

    n = (
        db.query(Analysis)
        .filter(
            Analysis.user_id == row.user_id,
            Analysis.contract_type == row.contract_type,
            Analysis.created_at <= row.created_at,
        )
        .count()
    )
    return f"{base}-{n:02d}"

def _gen_contract_display_name(db, user_id: int, contract_type: str) -> str:
    """
    为某个用户 + 合同类型生成稳定的展示名：
    通用合同-01 / 通用合同-02 / ...
    """
    base = CONTRACT_TYPE_CN.get(contract_type, "合同")

    count = (
        db.query(Analysis)
        .filter(
            Analysis.user_id == user_id,
            Analysis.contract_type == contract_type,
        )
        .count()
    )

    index = count + 1
    return f"{base}-{index:02d}"

def _safe_json_load(text: str) -> dict:
    t = (text or "").strip()

    # strip ```json ... ```
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z]*\n?", "", t)
        t = re.sub(r"\n?```$", "", t)
        t = t.strip()

    t = _CONTROL_CHARS_RE.sub("", t)
    return json.loads(t)


def _make_token(user_id: int) -> str:
    payload = {
        "sub": str(user_id),
        "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS),
        "iat": datetime.utcnow(),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALG)


def _get_user_id_from_auth(authorization: Optional[str]) -> int:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.split(" ", 1)[1].strip()
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        return int(payload["sub"])
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

def _wechat_jscode2session(code: str) -> Dict[str, Any]:
    if not WECHAT_APPID or not WECHAT_SECRET:
        raise HTTPException(status_code=500, detail="Missing WECHAT_APPID/WECHAT_SECRET on server")

    code = (code or "").strip()
    if not code:
        raise HTTPException(status_code=400, detail="Missing code")

    params = {
        "appid": WECHAT_APPID,
        "secret": WECHAT_SECRET,
        "js_code": code,
        "grant_type": "authorization_code",
    }
    url = "https://api.weixin.qq.com/sns/jscode2session?" + urllib.parse.urlencode(params)

    try:
        with urllib.request.urlopen(url, timeout=10) as resp:
            raw = resp.read().decode("utf-8")
            data = json.loads(raw)
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"WeChat jscode2session request failed: {e}")

    # WeChat 返回错误时会带 errcode/errmsg
    if isinstance(data, dict) and data.get("errcode"):
        raise HTTPException(status_code=401, detail=f"WeChat login failed: {data.get('errmsg', '')} ({data.get('errcode')})")

    return data

from demo_seed import DEMO_ANALYSES
def _seed_demo_analyses_if_needed(db, user_id: int):
    # 只在用户没有任何历史时种 demo（避免重复）
    existing = db.query(Analysis).filter(Analysis.user_id == user_id).count()
    if existing > 0:
        return

    for d in DEMO_ANALYSES[:2]:
        row = Analysis(
            user_id=user_id,
            contract_type=d["contract_type"],
            identity=d["identity"],
            prompt_version=PROMPT_VERSION,
            content_hash="demo_" + hashlib.sha256((d["display_name"]).encode("utf-8")).hexdigest(),
            original_content=d["original_content"],
            result_json=json.dumps(d["result_json"], ensure_ascii=False),
            file_name=None,
            file_mime=None,
            file_bytes=None,
            display_name=d["display_name"],
        )
        db.add(row)

    db.commit()

def _ensure_wechat_user(db, openid: str) -> User:
    # 用 phone 字段承载微信用户主键，避免改表
    wx_phone = f"wx_{openid}"

    u = db.query(User).filter(User.phone == wx_phone).first()
    if u:
        return u

    # password_hash 必填：给一个随机 hash（不会被使用）
    dummy_pw = hashlib.sha256(os.urandom(16)).hexdigest()[:6]  # 6 chars
    u = User(
        phone=wx_phone,
        password_hash=pwd_context.hash(dummy_pw),
        display_name="微信用户",
        avatar_url=None,
        credits=1
    )
    db.add(u)
    db.commit()
    db.refresh(u)
    _seed_demo_analyses_if_needed(db, u.id)
    return u

def _analyze_text_once(extracted_text: str, contract_type: str, identity: str) -> Dict[str, Any]:
    """
    Single-pass analysis call.
    """
    system = build_system_prompt(contract_type, identity)
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": extracted_text},
        ],
    )
    raw = (resp.choices[0].message.content or "").strip()
    parsed = _safe_json_load(raw)

    merged = dict(OUTPUT_TEMPLATE)
    if isinstance(parsed, dict):
        merged.update(parsed)
    if not merged.get("originalContent"):
        merged["originalContent"] = extracted_text
    return merged


def _merge_chunk_results(
    chunk_results: List[Dict[str, Any]],
    contract_type: str,
    identity: str,
) -> Dict[str, Any]:
    """
    Merge multiple chunk analyses into a single JSON using a final model pass.
    Keeps output schema strict.
    """
    # Prepare an input that the model can reliably merge.
    # Keep it compact: only score/riskSummary/clauses per chunk.
    compact = []
    for i, r in enumerate(chunk_results, start=1):
        compact.append(
            {
                "chunkIndex": i,
                "score": r.get("score", 0),
                "riskSummary": r.get("riskSummary", ""),
                "clauses": r.get("clauses", []) or [],
            }
        )

    merge_system = (
        build_system_prompt(contract_type, identity)
        + "\n你将收到多个分块审阅结果，请合并去重并输出最终JSON。"
        + "合并规则：\n"
          "- clauses 去重：若 originalText/标题/含义高度相似则只保留更清晰的一条。\n"
          "- 保留最重要的风险点，最多12条。\n"
          "- score 取整体风险评估（不是简单平均，可偏向更高风险）。\n"
    )

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        temperature=0,
        messages=[
            {"role": "system", "content": merge_system},
            {"role": "user", "content": json.dumps(compact, ensure_ascii=False)},
        ],
    )
    raw = (resp.choices[0].message.content or "").strip()
    parsed = _safe_json_load(raw)

    merged = dict(OUTPUT_TEMPLATE)
    if isinstance(parsed, dict):
        merged.update(parsed)
    return merged


def _analyze_with_chunking(
    full_text: str,
    contract_type: str,
    identity: str,
) -> Dict[str, Any]:
    """
    If text <= 80k chars -> single pass.
    If > 80k -> chunk into ~20k parts (max 12 chunks), analyze each, then merge.
    """
    text = (full_text or "").strip()
    if not text:
        return dict(OUTPUT_TEMPLATE)

    if len(text) <= MAX_TEXT_CHARS:
        out = _analyze_text_once(text, contract_type, identity)
        # Ensure originalContent present
        if not out.get("originalContent"):
            out["originalContent"] = text
        return out

    # Chunking
    chunks: List[str] = []
    start = 0
    while start < len(text) and len(chunks) < MAX_CHUNKS:
        chunks.append(text[start:start + CHUNK_CHARS])
        start += CHUNK_CHARS

    chunk_results: List[Dict[str, Any]] = []
    for idx, chunk in enumerate(chunks, start=1):
        # Add a small header so model understands chunk context
        chunk_input = f"【分块 {idx}/{len(chunks)}】\n{chunk}"
        r = _analyze_text_once(chunk_input, contract_type, identity)
        chunk_results.append(r)

    final = _merge_chunk_results(chunk_results, contract_type, identity)
    # Preserve original full content for UI
    final["originalContent"] = text
    return final

def _ocr_images_with_openai_single(image_bytes_list: List[bytes], mime_list: List[str]) -> str:
    """
    单次 OCR（最多 5 张图），只做一次 OpenAI 调用。
    """
    max_n = 5
    image_bytes_list = image_bytes_list[:max_n]
    mime_list = mime_list[:max_n]

    parts: List[Dict[str, Any]] = [
        {"type": "text", "text": "请对图片进行OCR，输出完整可读的中文合同文本。只输出纯文本，不要解释。"}
    ]

    for b, mime in zip(image_bytes_list, mime_list):
        b64 = __import__("base64").b64encode(b).decode("utf-8")
        parts.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}})

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[{"role": "user", "content": parts}],
    )
    return (resp.choices[0].message.content or "").strip()


def _ocr_images_with_openai_multi(image_bytes_list: List[bytes], mime_list: List[str]) -> str:
    """
    支持 1~9 张图：分批（每批最多 5 张）做 OCR，然后合并文本。
    """
    texts: List[str] = []
    batch_size = 5

    for i in range(0, len(image_bytes_list), batch_size):
        sub_bytes = image_bytes_list[i:i + batch_size]
        sub_mimes = mime_list[i:i + batch_size]
        t = _ocr_images_with_openai_single(sub_bytes, sub_mimes)
        if t:
            texts.append(t.strip())

    return "\n\n".join(texts).strip()


# ✅ 兼容旧调用：如果你代码里还有 _ocr_images_with_openai(...)，让它默认走 single
def _ocr_images_with_openai(image_bytes_list: List[bytes], mime_list: List[str]) -> str:
    return _ocr_images_with_openai_single(image_bytes_list, mime_list)

    parts: List[Dict[str, Any]] = [{"type": "text", "text": "请对图片进行OCR，输出完整可读的中文合同文本。只输出纯文本，不要解释。"}]
    for b, mime in zip(image_bytes_list, mime_list):
        b64 = __import__("base64").b64encode(b).decode("utf-8")
        parts.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}})

    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0,
        messages=[{"role": "user", "content": parts}],
    )
    return (resp.choices[0].message.content or "").strip()


def _extract_pdf_text_or_ocr(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    # 1) Extract selectable text (first 10 pages)
    texts = []
    for i in range(min(doc.page_count, 10)):
        t = doc.load_page(i).get_text("text").strip()
        if t:
            texts.append(t)
    extracted = "\n\n".join(texts).strip()
    if len(extracted) >= 200:
        return extracted

    # 2) OCR: render first 3 pages
    images: List[bytes] = []
    mimes: List[str] = []
    for i in range(min(doc.page_count, 3)):
        page = doc.load_page(i)
        pix = page.get_pixmap(dpi=200)
        images.append(pix.tobytes("png"))
        mimes.append("image/png")

    if not images:
        return extracted
    return _ocr_images_with_openai(images, mimes)

def _extract_docx_text(docx_bytes: bytes) -> str:
    f = io.BytesIO(docx_bytes)
    doc = Document(f)

    parts = []

    # 段落
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # 表格（合同关键信息经常在表格里）
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    cells.append(ct)
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()

@app.post("/v1/auth/register")
def register(req: RegisterRequest):
    phone, password = _validate_phone_password(req.phone, req.password)
    
    db = SessionLocal()
    try:
        if db.query(User).filter(User.phone == phone).first():
            raise HTTPException(status_code=409, detail="手机号已经被注册")
        u = User(phone=phone, password_hash=pwd_context.hash(password), display_name=f"用户{phone[-4:]}", avatar_url=None, credits=1)
        db.add(u)
        db.commit()
        return {"ok": True}
    finally:
        db.close()


@app.post("/v1/auth/login", response_model=LoginResponse)
def login(req: LoginRequest):
    phone, password = _validate_phone_password(req.phone, req.password)

    db = SessionLocal()
    try:
        u = db.query(User).filter(User.phone == phone).first()
        if not u or not pwd_context.verify(password, u.password_hash):
            raise HTTPException(status_code=401, detail="手机号或密码错误")
        return {"access_token": _make_token(u.id), "token_type": "bearer"}
    finally:
        db.close()

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

@app.post("/v1/auth/wechat/login", response_model=LoginResponse)
def wechat_login(req: WechatLoginRequest):
    # 1) code -> openid
    data = _wechat_jscode2session(req.code)
    openid = (data.get("openid") or "").strip()
    if not openid:
        raise HTTPException(status_code=401, detail="WeChat login failed: missing openid")

    # 2) find or create local user
    db = SessionLocal()
    try:
        u = _ensure_wechat_user(db, openid)
        # 3) issue our JWT (keep same mechanism as phone login)
        return {"access_token": _make_token(u.id), "token_type": "bearer"}
    finally:
        db.close()

@app.post("/v1/auth/change-password")
def change_password(req: ChangePasswordRequest, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)

    old_pw = (req.old_password or "").strip()
    new_pw = (req.new_password or "").strip()

    # 与前端一致：6位数字
    if not (len(old_pw) == 6 and old_pw.isdigit()):
        raise HTTPException(status_code=400, detail="旧密码应为 6 位数字")
    if not (len(new_pw) == 6 and new_pw.isdigit()):
        raise HTTPException(status_code=400, detail="新密码应为 6 位数字")
    if new_pw == old_pw:
        raise HTTPException(status_code=400, detail="新密码不能与旧密码相同")

    db = SessionLocal()
    try:
        u = db.query(User).filter(User.id == user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")

        if not pwd_context.verify(old_pw, u.password_hash):
            raise HTTPException(status_code=401, detail="旧密码不正确")

        u.password_hash = pwd_context.hash(new_pw)
        db.add(u)
        db.commit()

        # 前端只判断 res.ok，所以返回什么都行
        return {"ok": True}
    finally:
        db.close()

# =========================
# Credits Helpers (NEW)
# =========================

def _insufficient_credits(detail: str = "Insufficient credits", credits: int = 0):
    # 402 Payment Required
    raise HTTPException(
        status_code=402,
        detail={
            "message": detail,
            "code": "INSUFFICIENT_CREDITS",
            "credits": int(credits or 0),
        },
    )

def _require_credits(db, user_id: int, cost: int = 1) -> int:
    """
    Deduct credits before expensive analysis.
    Returns remaining credits after deduction.
    """
    u = db.query(User).filter(User.id == user_id).first()
    if not u:
        raise HTTPException(status_code=404, detail="User not found")

    current = int(u.credits or 0)
    if current < cost:
        _insufficient_credits(credits=current)

    u.credits = current - cost
    db.add(u)
    db.commit()
    db.refresh(u)
    return int(u.credits or 0)


# =========================
# Share Helpers (NEW)
# =========================
def _score_to_title(score: int) -> str:
    """
    Share 页用的标题，后端保证一致性（不要依赖前端）。
    """
    try:
        s = int(round(float(score)))
    except Exception:
        s = 0

    if s >= 90:
        return "整体安全"
    if s >= 70:
        return "整体安全"
    if s >= 40:
        return "检测到中度风险"
    return "检测到高风险"


def _gen_share_id() -> str:
    """
    不可预测的 share id：s_ + 12~16 位随机 hash 前缀
    """
    raw = hashlib.sha256(os.urandom(32)).hexdigest()
    return "s_" + raw[:16]

class MeResponse(BaseModel):
    id: int
    phone: str
    created_at: Optional[str] = None
    display_name: Optional[str] = None
    avatar_url: Optional[str] = None
    credits: int = 0

class UpdateMeRequest(BaseModel):
    display_name: Optional[str] = None
    avatar_url: Optional[str] = None


@app.get("/v1/users/me", response_model=MeResponse)
def get_me(authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    db = SessionLocal()
    try:
        u = db.query(User).filter(User.id == user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")
        return {
            "id": u.id,
            "phone": u.phone,
            "created_at": u.created_at.isoformat() if u.created_at else None,
            "display_name": u.display_name,
            "avatar_url": u.avatar_url,
            "credits": int(u.credits or 0),
        }
    finally:
        db.close()


@app.patch("/v1/users/me", response_model=MeResponse)
def patch_me(req: UpdateMeRequest, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    db = SessionLocal()
    try:
        u = db.query(User).filter(User.id == user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")

        if req.display_name is not None:
            u.display_name = req.display_name.strip() or u.display_name

        if req.avatar_url is not None:
            u.avatar_url = req.avatar_url.strip() or None

        db.add(u)
        db.commit()
        db.refresh(u)
        return {
            "id": u.id,
            "phone": u.phone,
            "created_at": u.created_at.isoformat() if u.created_at else None,
            "display_name": u.display_name,
            "avatar_url": u.avatar_url,
            "credits": int(u.credits or 0),
        }
    finally:
        db.close()

@app.put("/v1/users/me", response_model=MeResponse)
def put_me(req: UpdateMeRequest, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    db = SessionLocal()
    try:
        u = db.query(User).filter(User.id == user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")

        if req.display_name is not None:
            u.display_name = req.display_name.strip() or u.display_name

        if req.avatar_url is not None:
            u.avatar_url = req.avatar_url.strip() or None

        db.add(u)
        db.commit()
        db.refresh(u)
        return {
            "id": u.id,
            "phone": u.phone,
            "created_at": u.created_at.isoformat() if u.created_at else None,
            "display_name": u.display_name,
            "avatar_url": u.avatar_url,
            "credits": int(u.credits or 0),
        }
    finally:
        db.close()

@app.post("/v1/users/me/avatar")
async def upload_my_avatar(
    authorization: Optional[str] = Header(default=None),
    file: UploadFile = File(...),
):
    user_id = _get_user_id_from_auth(authorization)

    # 1) 校验文件类型
    ctype = (file.content_type or "").lower()
    if not ctype.startswith("image/"):
        raise HTTPException(status_code=400, detail="Only image files are allowed")

    # 2) 读取 bytes + 简单限大小（建议 3MB，够头像了）
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(data) > 3 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="Avatar too large (max 3MB)")

    # 3) 生成文件名
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in [".jpg", ".jpeg", ".png", ".webp"]:
        # 如果没扩展名或奇怪扩展名，按 content-type 兜底
        if ctype == "image/png":
            ext = ".png"
        elif ctype == "image/webp":
            ext = ".webp"
        else:
            ext = ".jpg"

    fname = f"{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(AVATAR_DIR, fname)

    # 4) 写入磁盘
    with open(save_path, "wb") as f:
        f.write(data)

    # 5) 写回 DB（存相对路径，未来换域名更方便）
    avatar_url = f"/uploads/avatars/{fname}"

    db = SessionLocal()
    try:
        u = db.query(User).filter(User.id == user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")

        u.avatar_url = avatar_url
        db.add(u)
        db.commit()
        db.refresh(u)

        return {"avatar_url": u.avatar_url}
    finally:
        db.close()

@app.post("/v1/contracts/analyze/text", response_model=AnalysisResult)
def analyze_text(req: AnalyzeTextRequest, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)

    t = validate_contract_type(req.type)
    content = (req.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="Empty content")

    cache_key_raw = (PROMPT_VERSION + "|" + t + "|" + req.identity + "|" + content).encode("utf-8")
    content_hash = hashlib.sha256(cache_key_raw).hexdigest()

    db = SessionLocal()
    try:
        cached = (
            db.query(Analysis)
            .filter(
                Analysis.user_id == user_id,
                Analysis.contract_type == t,
                Analysis.identity == req.identity,
                Analysis.prompt_version == PROMPT_VERSION,
                Analysis.content_hash == content_hash,
            )
            .order_by(Analysis.created_at.desc())
            .first()
        )
        if cached:
            data = json.loads(cached.result_json)
            row_id = cached.id
            display_name = (cached.display_name or "").strip()
            if not display_name:
                display_name = _gen_contract_display_name(db, user_id, t)
                cached.display_name = display_name
                db.add(cached)
                db.commit()
        else:
            _require_credits(db, user_id, cost=1)
            data = _analyze_with_chunking(content, t, req.identity)
            display_name = _gen_contract_display_name(db, user_id, t)
            new_row = Analysis(
                user_id=user_id,
                contract_type=t,
                identity=req.identity,
                prompt_version=PROMPT_VERSION,
                content_hash=content_hash,
                original_content=content,
                result_json=json.dumps(data, ensure_ascii=False),
                file_name=None,
                file_mime=None,
                file_bytes=None,
                display_name=display_name,
            )
            db.add(new_row)
            db.commit()
            db.refresh(new_row)
            row_id = new_row.id

        now = datetime.now().isoformat()
        display_name = _gen_contract_display_name(db, user_id, t)

        return {
            "id": str(row_id),
            "name": display_name,
            "date": now,
            "score": float(data.get("score", 0) or 0),
            "riskSummary": str(data.get("riskSummary", "") or ""),
            "clauses": data.get("clauses", []) or [],
            "originalContent": str(data.get("originalContent", content) or content),
            "status": "completed",
            "type": t,
            "identity": req.identity,
            "promptVersion": PROMPT_VERSION,
            "imagePreview": None,
            "fileUrl": None,  # TEXT analysis has no original file to download
        }
    finally:
        db.close()


@app.post("/v1/contracts/analyze/upload", response_model=AnalysisResult)
async def analyze_upload(
    authorization: Optional[str] = Header(default=None),
    type: str = Form("general"),
    identity: Literal["A", "B"] = Form(...),
    file: UploadFile = File(...),
):
    user_id = _get_user_id_from_auth(authorization)

    t = validate_contract_type(type)

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty file")

    # Hard limit: 50MB
    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 50MB)")

    mime = file.content_type or "application/octet-stream"
    filename = file.filename or "upload"

    # cache based on raw file hash + prompt version + identity/type
    file_sha = hashlib.sha256(file_bytes).hexdigest()
    cache_key_raw = (PROMPT_VERSION + "|" + t + "|" + identity + "|" + file_sha).encode("utf-8")
    content_hash = hashlib.sha256(cache_key_raw).hexdigest()

    db = SessionLocal()
    try:
        cached = (
            db.query(Analysis)
            .filter(
                Analysis.user_id == user_id,
                Analysis.contract_type == t,
                Analysis.identity == identity,
                Analysis.prompt_version == PROMPT_VERSION,
                Analysis.content_hash == content_hash,
            )
            .order_by(Analysis.created_at.desc())
            .first()
        )
        if cached:
            data = json.loads(cached.result_json)
            row_id = cached.id
            display_name = (cached.display_name or "").strip()
            if not display_name:
                display_name = _gen_contract_display_name(db, user_id, t)
                cached.display_name = display_name
                db.add(cached)
                db.commit()
        else:
            _require_credits(db, user_id, cost=1)
            # 1) extract text
            ext = (os.path.splitext(filename)[1] or "").lower()
            mime = (file.content_type or "application/octet-stream").lower()
            
            # ✅ 新增：docx 先抽文本（不 OCR、不转 PDF）
            if ext == ".docx" or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                extracted_text = _extract_docx_text(file_bytes)

            elif ext == ".pdf" or mime == "application/pdf":
                extracted_text = _extract_pdf_text_or_ocr(file_bytes)

            # （如果你要保留图片，就加上 image 分支；你现在说只要 pdf/docx，就不加）
            elif mime.startswith("image/") or ext in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
              # mime 兜底：如果 mime 不像 image，就按扩展名推断
                if not mime.startswith("image/"):
                    if ext in [".jpg", ".jpeg"]:
                        mime = "image/jpeg"
                    elif ext == ".webp":
                        mime = "image/webp"
                    else:
                        mime = "image/png"

                extracted_text = _ocr_images_with_openai([file_bytes], [mime])
            
            else:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {mime} ({ext})")

            extracted_text = (extracted_text or "").strip() or "（解析失败：未提取到文本）"

            # 2) analyze (with chunking if needed)
            data = _analyze_with_chunking(extracted_text, t, identity)
            display_name = _gen_contract_display_name(db, user_id, t)
            new_row = Analysis(
                user_id=user_id,
                contract_type=t,
                identity=identity,
                prompt_version=PROMPT_VERSION,
                content_hash=content_hash,
                original_content=extracted_text,
                result_json=json.dumps(data, ensure_ascii=False),
                file_name=filename,
                file_mime=mime,
                file_bytes=file_bytes,
                display_name=display_name,
            )
            db.add(new_row)
            db.commit()
            db.refresh(new_row)
            row_id = new_row.id

        now = datetime.now().isoformat()

        return {
            "id": str(row_id),
            "name": display_name,
            "date": now,
            "score": float(data.get("score", 0) or 0),
            "riskSummary": str(data.get("riskSummary", "") or ""),
            "clauses": data.get("clauses", []) or [],
            "originalContent": str(data.get("originalContent", "") or ""),
            "status": "completed",
            "type": t,
            "identity": identity,
            "promptVersion": PROMPT_VERSION,
            "imagePreview": None,
            "fileUrl": f"/v1/contracts/{row_id}/file",
            "fileName": filename,
        }
    finally:
        db.close()

@app.post("/v1/contracts/analyze/upload/batch", response_model=AnalysisResult)
async def analyze_upload_batch(
    authorization: Optional[str] = Header(default=None),
    batch_id: str = Form(...),
    idx: int = Form(...),     # 1..N
    total: int = Form(...),   # N
    type: str = Form("general"),
    identity: Literal["A", "B"] = Form(...),
    file: UploadFile = File(...),
):
    user_id = _get_user_id_from_auth(authorization)

    batch_id = (batch_id or "").strip()
    if not batch_id:
        raise HTTPException(status_code=400, detail="Missing batch_id")

    if total < 1 or total > 9:
        raise HTTPException(status_code=400, detail="total must be 1..9")
    if idx < 1 or idx > total:
        raise HTTPException(status_code=400, detail="idx out of range")

    t = validate_contract_type(type)

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(file_bytes) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 50MB)")

    mime = (file.content_type or "application/octet-stream").lower()
    filename = file.filename or f"page_{idx}"

    # 只允许图片走 batch
    ext = (os.path.splitext(filename)[1] or "").lower()
    if not (mime.startswith("image/") or ext in [".png", ".jpg", ".jpeg", ".webp"]):
        raise HTTPException(status_code=400, detail=f"Batch endpoint supports images only: {mime} ({ext})")

    # mime 兜底
    if not mime.startswith("image/"):
        if ext in [".jpg", ".jpeg"]:
            mime = "image/jpeg"
        elif ext == ".webp":
            mime = "image/webp"
        else:
            mime = "image/png"

    db = SessionLocal()
    try:
        # upsert batch
        batch = (
            db.query(UploadBatch)
            .filter(UploadBatch.user_id == user_id, UploadBatch.batch_id == batch_id)
            .first()
        )
        if not batch:
            batch = UploadBatch(
                user_id=user_id,
                batch_id=batch_id,
                contract_type=t,
                identity=identity,
                total=total,
                received=0,
            )
            db.add(batch)
            db.commit()
            db.refresh(batch)
        else:
            # 防御：同 batch_id 不能换参数
            if batch.contract_type != t or batch.identity != identity or batch.total != total:
                raise HTTPException(status_code=400, detail="Batch params mismatch")

        # idx 覆盖写入（允许重传）
        existing = (
            db.query(UploadBatchFile)
            .filter(
                UploadBatchFile.user_id == user_id,
                UploadBatchFile.batch_id == batch_id,
                UploadBatchFile.idx == idx,
            )
            .first()
        )
        if existing:
            existing.file_name = filename
            existing.file_mime = mime
            existing.file_bytes = file_bytes
            db.add(existing)
        else:
            db.add(
                UploadBatchFile(
                    user_id=user_id,
                    batch_id=batch_id,
                    idx=idx,
                    file_name=filename,
                    file_mime=mime,
                    file_bytes=file_bytes,
                )
            )
        db.commit()

        # 重新统计 received
        received = (
            db.query(UploadBatchFile)
            .filter(UploadBatchFile.user_id == user_id, UploadBatchFile.batch_id == batch_id)
            .count()
        )
        batch.received = received
        db.add(batch)
        db.commit()

        # ✅ 永远快速返回：让前端收齐后再调 finalize
        now = datetime.now().isoformat()
        display_name = _gen_contract_display_name(db, user_id, t)
        return {
            "id": "0",
            "name": display_name,
            "date": now,
            "score": 0.0,
            "riskSummary": "",
            "clauses": [],
            "originalContent": "",
            "status": "completed",
            "type": t,
            "identity": identity,
            "promptVersion": PROMPT_VERSION,
            "imagePreview": None,
            "fileUrl": None,
            "meta": {
                "processedImages": received,
                "totalImages": total,
                "batchId": batch_id,
                "pending": True,
                "readyToFinalize": (received >= total),
            },
        }
    finally:
        db.close()

@app.post("/v1/contracts/analyze/upload/batch/finalize", response_model=AnalysisResult)
def finalize_upload_batch(req: BatchFinalizeRequest, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    batch_id = (req.batch_id or "").strip()
    if not batch_id:
        raise HTTPException(status_code=400, detail="Missing batch_id")

    db = SessionLocal()
    try:
        batch = (
            db.query(UploadBatch)
            .filter(UploadBatch.user_id == user_id, UploadBatch.batch_id == batch_id)
            .first()
        )
        if not batch:
            raise HTTPException(status_code=404, detail="Batch not found")

        total = batch.total
        received = (
            db.query(UploadBatchFile)
            .filter(UploadBatchFile.user_id == user_id, UploadBatchFile.batch_id == batch_id)
            .count()
        )
        if received < total:
            raise HTTPException(status_code=400, detail=f"Batch not complete: received {received}/{total}")

        _require_credits(db, user_id, cost=1)

        rows = (
            db.query(UploadBatchFile)
            .filter(UploadBatchFile.user_id == user_id, UploadBatchFile.batch_id == batch_id)
            .order_by(UploadBatchFile.idx.asc())
            .all()
        )
        image_bytes_list = [r.file_bytes for r in rows]
        mime_list = [(r.file_mime or "image/png") for r in rows]

        extracted_text = _ocr_images_with_openai_multi(image_bytes_list, mime_list)
        extracted_text = (extracted_text or "").strip() or "（解析失败：未提取到文本）"

        t = batch.contract_type
        identity = batch.identity

        data = _analyze_with_chunking(extracted_text, t, identity)

        cache_key_raw = (PROMPT_VERSION + "|" + t + "|" + identity + "|" + hashlib.sha256(extracted_text.encode("utf-8")).hexdigest()).encode("utf-8")
        content_hash = hashlib.sha256(cache_key_raw).hexdigest()
        display_name= _gen_contract_display_name(db, user_id, t)
        new_row = Analysis(
            user_id=user_id,
            contract_type=t,
            identity=identity,
            prompt_version=PROMPT_VERSION,
            content_hash=content_hash,
            original_content=extracted_text,
            result_json=json.dumps(data, ensure_ascii=False),
            file_name=None,
            file_mime=None,
            file_bytes=None,
            display_name=display_name,
        )
        db.add(new_row)
        db.commit()
        db.refresh(new_row)

        # cleanup
        db.query(UploadBatchFile).filter(
            UploadBatchFile.user_id == user_id,
            UploadBatchFile.batch_id == batch_id
        ).delete(synchronize_session=False)
        db.query(UploadBatch).filter(
            UploadBatch.user_id == user_id,
            UploadBatch.batch_id == batch_id
        ).delete(synchronize_session=False)
        db.commit()

        now = datetime.now().isoformat()
        return {
            "id": str(new_row.id),
            "name": display_name,
            "date": now,
            "score": float(data.get("score", 0) or 0),
            "riskSummary": str(data.get("riskSummary", "") or ""),
            "clauses": data.get("clauses", []) or [],
            "originalContent": str(data.get("originalContent", extracted_text) or extracted_text),
            "status": "completed",
            "type": t,
            "identity": identity,
            "promptVersion": PROMPT_VERSION,
            "imagePreview": None,
            "fileUrl": None,
            "meta": {"processedImages": total, "totalImages": total, "batchId": batch_id, "pending": False},
        }
    finally:
        db.close()

# =========================
# Share Endpoints (NEW)
# =========================

@app.post("/v1/shares", response_model=CreateShareResponse)
def create_share(req: CreateShareRequest, authorization: Optional[str] = Header(default=None)):
    """
    创建分享（30天有效）
    - 需要登录
    - contractName 一律冻结为 Analysis.display_name（单一真相源）
    - 不依赖前端传入 contractName（只做兼容）
    """
    user_id = _get_user_id_from_auth(authorization)

    try:
        analysis_id = int(str(req.analysis_id).strip())
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid analysis_id")

    db = SessionLocal()
    try:
        row = (
            db.query(Analysis)
            .filter(Analysis.id == analysis_id, Analysis.user_id == user_id)
            .first()
        )
        if not row:
            raise HTTPException(status_code=404, detail="Analysis not found")

        # 读分析结果
        try:
            data = json.loads(row.result_json or "{}")
        except Exception:
            data = {}

        score = data.get("score", 0) or 0
        try:
            score_int = int(round(float(score)))
        except Exception:
            score_int = 0

        risk_summary = str(data.get("riskSummary", "") or "").strip() or "暂无风险摘要"
        score_title = _score_to_title(score_int)

        # ✅ contractName：唯一来源 = Analysis.display_name
        # 兼容：如果历史数据没有 display_name，就用 _analysis_display_name(db, row) 现场算一个
        display_name = (getattr(row, "display_name", None) or "").strip()
        if not display_name:
            display_name = _analysis_display_name(db, row)  # 你已有这个函数

        # 生成 share_id（确保唯一）
        share_id = _gen_share_id()
        for _ in range(3):
            exists = db.query(AnalysisShare).filter(AnalysisShare.share_id == share_id).first()
            if not exists:
                break
            share_id = _gen_share_id()
        else:
            raise HTTPException(status_code=500, detail="Failed to generate shareId")

        expires_at = datetime.utcnow() + timedelta(days=30)

        # 写入分享快照（冻结 display_name）
        srow = AnalysisShare(
            share_id=share_id,
            analysis_id=analysis_id,
            user_id=user_id,
            score=score_int,
            score_title=score_title,
            risk_summary=risk_summary,
            contract_name=display_name,
            expires_at=expires_at,
        )
        db.add(srow)
        db.commit()

        return {"shareId": share_id, "expiresAt": expires_at.isoformat()}
    finally:
        db.close()

@app.get("/v1/shares/{share_id}", response_model=SharePublicResponse)
def get_share_public(share_id: str):
    """
    获取分享摘要（公开接口，无需登录）
    - 过期：返回 410 Gone（更语义化）
    """
    sid = (share_id or "").strip()
    if not sid:
        raise HTTPException(status_code=404, detail="Share not found")

    db = SessionLocal()
    try:
        row = db.query(AnalysisShare).filter(AnalysisShare.share_id == sid).first()
        if not row:
            raise HTTPException(status_code=404, detail="Share not found")

        # 过期判断
        now = datetime.utcnow()
        if row.expires_at and row.expires_at < now:
            # 410：资源已失效（前端可显示“分享已过期”）
            raise HTTPException(status_code=410, detail="Share expired")

        return {
  	    "contractName": row.contract_name or "合同风险分析",
            "score": int(row.score or 0),
            "scoreTitle": row.score_title or _score_to_title(int(row.score or 0)),
            "riskSummary": row.risk_summary or "暂无风险摘要",
            "expiresAt": row.expires_at.isoformat() if row.expires_at else "",
        }
    finally:
        db.close()

@app.get("/v1/contracts/history")
def history(authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    db = SessionLocal()
    try:
        rows = (
            db.query(Analysis)
            .filter(Analysis.user_id == user_id)
            .order_by(Analysis.created_at.desc())
            .limit(50)
            .all()
        )
        return {
            "items": [
                {
                    "id": str(r.id),
                    "type": r.contract_type,
                    "identity": r.identity,
                    "promptVersion": r.prompt_version,
                    "date": r.created_at.isoformat(),
                    "result": json.loads(r.result_json),
                    "fileUrl": f"/v1/contracts/{r.id}/file" if r.file_bytes else None,
		    "fileName": r.file_name,
                }
                for r in rows
            ]
        }
    finally:
        db.close()


@app.get("/v1/contracts/{analysis_id}/file")
def download_file(analysis_id: int, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    db = SessionLocal()
    try:
        row = (
            db.query(Analysis)
            .filter(Analysis.id == analysis_id, Analysis.user_id == user_id)
            .first()
        )
        if not row or not row.file_bytes:
            raise HTTPException(status_code=404, detail="File not found")

        mime = row.file_mime or "application/octet-stream"
        filename = row.file_name or f"analysis_{analysis_id}"

        return Response(
            content=row.file_bytes,
            media_type=mime,
            headers={"Content-Disposition": f'attachment; filename=\"{filename}\"'},
        )
    finally:
        db.close()

@app.delete("/v1/analyses/{analysis_id}")
def delete_one_analysis(analysis_id: int, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)

    db = SessionLocal()
    try:
        row = (
            db.query(Analysis)
            .filter(Analysis.id == analysis_id, Analysis.user_id == user_id)
            .first()
        )
        if not row:
            raise HTTPException(status_code=404, detail="Analysis not found")

        db.delete(row)
        db.commit()

        return {"ok": True, "deleted": 1, "id": analysis_id}
    finally:
        db.close()

@app.delete("/v1/analyses")
def wipe_all_analyses(authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)

    db = SessionLocal()
    try:
        # 先数一下要删多少条
        count = db.query(Analysis).filter(Analysis.user_id == user_id).count()

        # 删除该用户所有分析记录（包含 file_bytes，所以“原文件”也没了）
        db.query(Analysis).filter(Analysis.user_id == user_id).delete(synchronize_session=False)
        db.commit()

        return {"deleted": count}
    finally:
        db.close()

# =========================
# Pay Orders (REAL v3 JSAPI + query to credit)
# =========================
import time
import base64
import secrets
from urllib import request as urlrequest
from urllib.error import HTTPError, URLError

from sqlalchemy import Boolean
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import padding

class PayOrder(Base):
    __tablename__ = "pay_orders"
    id = Column(Integer, primary_key=True)
    user_id = Column(Integer, index=True, nullable=False)

    out_trade_no = Column(String(64), unique=True, index=True, nullable=False)
    sku_id = Column(String(32), nullable=False)

    total_fee = Column(Integer, nullable=False, default=0)  # 分
    credits = Column(Integer, nullable=False, default=0)

    status = Column(String(16), nullable=False, default="CREATED")  # CREATED / PAID / CLOSED
    wx_prepay_id = Column(String(128), nullable=True)
    wx_transaction_id = Column(String(128), nullable=True)

    created_at = Column(DateTime, default=datetime.utcnow)
    paid_at = Column(DateTime, nullable=True)

def _migrate_pay_tables():
    Base.metadata.create_all(bind=engine)
    with engine.connect() as conn:
        conn.exec_driver_sql("CREATE UNIQUE INDEX IF NOT EXISTS ix_pay_orders_out_trade_no ON pay_orders (out_trade_no);")
        conn.exec_driver_sql("CREATE INDEX IF NOT EXISTS ix_pay_orders_user_id ON pay_orders (user_id);")

_migrate_pay_tables()

# 充值 SKU（先写死，后续可从 DB/配置读取）
SKU_CATALOG = {
    "CREDIT_10":  {"credits": 1,  "total_fee": 5},   # 9.90
    "CREDIT_30":  {"credits": 3,  "total_fee": 2490},  # 24.90
    "CREDIT_100": {"credits": 10, "total_fee": 6990},  # 69.90
}

class PrepayRequest(BaseModel):
    sku_id: str

class PrepayResponse(BaseModel):
    outTradeNo: str
    paymentParams: Dict[str, Any]

WECHATPAY_BASE = "https://api.mch.weixin.qq.com"

def _gen_out_trade_no(user_id: int) -> str:
    tail = hashlib.sha256(os.urandom(16)).hexdigest()[:10]
    return f"o_{user_id}_{tail}"

def _load_merchant_private_key(path: str):
    with open(path, "rb") as f:
        key_data = f.read()
    return serialization.load_pem_private_key(key_data, password=None)

def _sign_rsa_sha256_base64(private_key, message: str) -> str:
    sig = private_key.sign(
        message.encode("utf-8"),
        padding.PKCS1v15(),
        hashes.SHA256(),
    )
    return base64.b64encode(sig).decode("utf-8")

def _require_wechat_pay_env() -> Dict[str, str]:
    need = {
        "WECHAT_APPID": os.getenv("WECHAT_APPID", "").strip(),
        "WECHAT_MCHID": os.getenv("WECHAT_MCHID", "").strip(),
        "WECHAT_MCH_SERIAL": os.getenv("WECHAT_MCH_SERIAL", "").strip(),
        "WECHAT_MCH_PRIVATE_KEY_PATH": os.getenv("WECHAT_MCH_PRIVATE_KEY_PATH", "").strip(),
        "WECHAT_PAY_NOTIFY_URL": os.getenv("WECHAT_PAY_NOTIFY_URL", "").strip(),  # 先保留，但这一步不会用它入账
    }
    missing = [k for k, v in need.items() if not v]
    if missing:
        raise HTTPException(
            status_code=500,
            detail={"message": "WeChat Pay not configured", "missing": missing},
        )

    # path 可读性检查
    p = need["WECHAT_MCH_PRIVATE_KEY_PATH"]
    if not os.path.exists(p):
        raise HTTPException(status_code=500, detail={"message": "Private key not found", "path": p})
    if not os.path.isfile(p):
        raise HTTPException(status_code=500, detail={"message": "Private key path is not a file", "path": p})
    return need

def _wechat_v3_request(method: str, path_with_query: str, body_obj: Optional[dict] = None) -> dict:
    """
    WeChat Pay v3 signed request. No client cert needed for normal v3 calls.
    """
    env = _require_wechat_pay_env()
    appid = env["WECHAT_APPID"]
    mchid = env["WECHAT_MCHID"]
    serial_no = env["WECHAT_MCH_SERIAL"]
    key_path = env["WECHAT_MCH_PRIVATE_KEY_PATH"]

    private_key = _load_merchant_private_key(key_path)

    ts = str(int(time.time()))
    nonce = secrets.token_urlsafe(16)

    body_str = ""
    if body_obj is not None:
        body_str = json.dumps(body_obj, ensure_ascii=False, separators=(",", ":"))

    # 注意：签名串里的 URL 只包含 path + query，不含域名
    sign_message = f"{method}\n{path_with_query}\n{ts}\n{nonce}\n{body_str}\n"
    signature = _sign_rsa_sha256_base64(private_key, sign_message)

    auth = (
        f'WECHATPAY2-SHA256-RSA2048 mchid="{mchid}",'
        f'nonce_str="{nonce}",timestamp="{ts}",serial_no="{serial_no}",signature="{signature}"'
    )

    url = WECHATPAY_BASE + path_with_query
    headers = {
        "Accept": "application/json",
        "User-Agent": "safecontract/1.0",
        "Authorization": auth,
    }
    if body_obj is not None:
        headers["Content-Type"] = "application/json; charset=utf-8"

    data = body_str.encode("utf-8") if body_obj is not None else None
    req = urlrequest.Request(url=url, data=data, method=method, headers=headers)

    try:
        with urlrequest.urlopen(req, timeout=15) as resp:
            raw = resp.read().decode("utf-8")
            return json.loads(raw) if raw else {}
    except HTTPError as e:
        raw = ""
        try:
            raw = e.read().decode("utf-8")
        except Exception:
            pass
        # 把微信侧错误完整透出，便于你定位
        raise HTTPException(status_code=502, detail={"message": "WeChat Pay HTTPError", "status": e.code, "body": raw})
    except URLError as e:
        raise HTTPException(status_code=502, detail={"message": "WeChat Pay URLError", "error": str(e)})

def _extract_openid_from_user(u: User) -> str:
    """
    你现在用 phone 字段承载微信主键：wx_<openid>
    """
    phone = (u.phone or "").strip()
    if phone.startswith("wx_") and len(phone) > 3:
        return phone[3:]
    return ""

def _build_jsapi_payment_params(appid: str, mch_private_key, prepay_id: str) -> Dict[str, Any]:
    """
    返回给小程序 wx.requestPayment 的参数（字段名必须匹配微信要求）
    """
    timeStamp = str(int(time.time()))
    nonceStr = secrets.token_urlsafe(16)
    package = f"prepay_id={prepay_id}"
    signType = "RSA"

    # paySign 签名串：appId\ntimeStamp\nnonceStr\npackage\n
    pay_sign_message = f"{appid}\n{timeStamp}\n{nonceStr}\n{package}\n"
    paySign = _sign_rsa_sha256_base64(mch_private_key, pay_sign_message)

    return {
        "timeStamp": timeStamp,
        "nonceStr": nonceStr,
        "package": package,
        "signType": signType,
        "paySign": paySign,
    }

@app.post("/v1/pay/prepay", response_model=PrepayResponse)
def pay_prepay(req: PrepayRequest, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    sku_id = (req.sku_id or "").strip()
    if sku_id not in SKU_CATALOG:
        raise HTTPException(status_code=400, detail="Invalid sku_id")

    env = _require_wechat_pay_env()
    appid = env["WECHAT_APPID"]
    mchid = env["WECHAT_MCHID"]
    key_path = env["WECHAT_MCH_PRIVATE_KEY_PATH"]

    sku = SKU_CATALOG[sku_id]
    credits = int(sku["credits"])
    total_fee = int(sku["total_fee"])

    db = SessionLocal()
    try:
        u = db.query(User).filter(User.id == user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")

        openid = _extract_openid_from_user(u)
        if not openid:
            raise HTTPException(status_code=400, detail="This account is not a WeChat user (missing openid)")

        out_trade_no = _gen_out_trade_no(user_id)

        # 先创建本地订单
        row = PayOrder(
            user_id=user_id,
            out_trade_no=out_trade_no,
            sku_id=sku_id,
            total_fee=total_fee,
            credits=credits,
            status="CREATED",
        )
        db.add(row)
        db.commit()
        db.refresh(row)

        # 调微信 JSAPI 预下单
        body = {
            "appid": appid,
            "mchid": mchid,
            "description": f"签前助手充值 {credits} 芒果币",
            "out_trade_no": out_trade_no,
            # notify_url 仍然要给（微信要求），但我们这一步不依赖它入账
            "notify_url": env["WECHAT_PAY_NOTIFY_URL"],
            "amount": {"total": total_fee, "currency": "CNY"},
            "payer": {"openid": openid},
        }
        resp = _wechat_v3_request("POST", "/v3/pay/transactions/jsapi", body_obj=body)
        prepay_id = (resp.get("prepay_id") or "").strip()
        if not prepay_id:
            raise HTTPException(status_code=502, detail={"message": "WeChat prepay failed (no prepay_id)", "resp": resp})

        # 生成 requestPayment 参数
        mch_private_key = _load_merchant_private_key(key_path)
        payment_params = _build_jsapi_payment_params(appid, mch_private_key, prepay_id)

        row.wx_prepay_id = prepay_id
        db.add(row)
        db.commit()

        return {"outTradeNo": out_trade_no, "paymentParams": payment_params}
    finally:
        db.close()

def _credit_if_wechat_success(db, order: PayOrder):
    """
    若微信已支付成功，则入账（幂等）
    """
    if order.status == "PAID":
        return

    env = _require_wechat_pay_env()
    mchid = env["WECHAT_MCHID"]

    # 查微信订单状态
    path = f"/v3/pay/transactions/out-trade-no/{order.out_trade_no}?mchid={mchid}"
    resp = _wechat_v3_request("GET", path_with_query=path, body_obj=None)

    trade_state = (resp.get("trade_state") or "").strip()
    # SUCCESS 才入账
    if trade_state == "SUCCESS":
        # 幂等：再次检查
        if order.status == "PAID":
            return

        txid = (resp.get("transaction_id") or "").strip() or None

        # 给用户加 credits
        u = db.query(User).filter(User.id == order.user_id).first()
        if not u:
            raise HTTPException(status_code=404, detail="User not found")

        u.credits = int(u.credits or 0) + int(order.credits or 0)
        order.status = "PAID"
        order.wx_transaction_id = txid
        order.paid_at = datetime.utcnow()

        db.add(u)
        db.add(order)
        db.commit()

@app.get("/v1/pay/orders/{out_trade_no}")
def pay_get_order(out_trade_no: str, authorization: Optional[str] = Header(default=None)):
    user_id = _get_user_id_from_auth(authorization)
    no = (out_trade_no or "").strip()
    if not no:
        raise HTTPException(status_code=404, detail="Not found")

    db = SessionLocal()
    try:
        row = (
            db.query(PayOrder)
            .filter(PayOrder.out_trade_no == no, PayOrder.user_id == user_id)
            .first()
        )
        if not row:
            raise HTTPException(status_code=404, detail="Not found")

        # 如果还没 PAID，就尝试向微信查询一次，成功则入账
        if row.status != "PAID":
            try:
                _credit_if_wechat_success(db, row)
                # refresh
                db.refresh(row)
            except HTTPException:
                # 微信查询失败也不影响返回订单本身
                pass
            except Exception:
                pass

        return {
            "outTradeNo": row.out_trade_no,
            "status": row.status,
            "credits": int(row.credits or 0),
            "totalFee": int(row.total_fee or 0),
            "createdAt": row.created_at.isoformat() if row.created_at else None,
            "paidAt": row.paid_at.isoformat() if row.paid_at else None,
        }
    finally:
        db.close()

from typing import Optional

@app.get("/v1/pay/orders")
def pay_list_orders(
    authorization: Optional[str] = Header(default=None),
    limit: int = 20,
):
    user_id = _get_user_id_from_auth(authorization)

    # 防御：避免一次拉太多
    if limit < 1:
        limit = 1
    if limit > 50:
        limit = 50

    db = SessionLocal()
    try:
        rows = (
            db.query(PayOrder)
            .filter(PayOrder.user_id == user_id)
            .order_by(PayOrder.created_at.desc())
            .limit(limit)
            .all()
        )

        return {
            "items": [
                {
                    "outTradeNo": r.out_trade_no,
                    "skuId": r.sku_id,
                    "status": r.status,
                    "credits": int(r.credits or 0),
                    "totalFee": int(r.total_fee or 0),
                    "createdAt": r.created_at.isoformat() if r.created_at else None,
                    "paidAt": r.paid_at.isoformat() if r.paid_at else None,
                    "wxTransactionId": r.wx_transaction_id,
                }
                for r in rows
            ]
        }
    finally:
        db.close()



# =========================
# WeChat Pay Notify (v3) - Public Key Mode: verify signature + decrypt + credit
# =========================
from fastapi import Request
from fastapi.responses import JSONResponse
from cryptography.hazmat.primitives.ciphers.aead import AESGCM
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.asymmetric import padding as asy_padding
from cryptography.hazmat.primitives.serialization import load_pem_public_key


def _get_api_v3_key() -> str:
    return os.getenv("WECHAT_API_V3_KEY", "").strip()


def _get_wechatpay_public_key_path() -> str:
    return os.getenv("WECHAT_PAY_PUBLIC_KEY_PATH", "").strip()


def _get_wechatpay_public_key_id_expected() -> str:
    # 可选：对照回调头 Wechatpay-Serial，防止用错 key
    return os.getenv("WECHAT_PAY_PUBLIC_KEY_ID", "").strip()


def _aes256gcm_decrypt(api_v3_key: str, nonce: str, ciphertext: str, associated_data: str) -> str:
    key_bytes = api_v3_key.encode("utf-8")
    aesgcm = AESGCM(key_bytes)

    ct = base64.b64decode(ciphertext)
    ad = (associated_data or "").encode("utf-8")
    n = (nonce or "").encode("utf-8")

    pt = aesgcm.decrypt(n, ct, ad)
    return pt.decode("utf-8")



def _hget(headers: dict, name: str) -> str:
    """
    Starlette/FastAPI 的 request.headers 是大小写不敏感，
    但 dict(request.headers) 后 key 往往是小写。
    这里做一次大小写无关读取。
    """
    if not headers:
        return ""
    # 先直接取
    v = headers.get(name)
    if v:
        return str(v).strip()
    # 再按小写取
    v = headers.get(name.lower())
    if v:
        return str(v).strip()
    # 最后遍历一次兜底（最稳）
    nl = name.lower()
    for k, vv in headers.items():
        if str(k).lower() == nl:
            return str(vv).strip()
    return ""


def _verify_wechatpay_signature_public_key_mode(
    *,
    raw_body: bytes,
    headers: dict,
) -> bool:
    """
    公钥模式验签（大小写无关 header 读取 + 更稳的 body 处理）
    """
    ts = _hget(headers, "Wechatpay-Timestamp")
    nonce = _hget(headers, "Wechatpay-Nonce")
    signature_b64 = _hget(headers, "Wechatpay-Signature")
    serial = _hget(headers, "Wechatpay-Serial")

    if not (ts and nonce and signature_b64 and serial):
        return False

    # 可选：对照公钥ID（PUB_KEY_ID...）
    expected_id = _get_wechatpay_public_key_id_expected()
    # ⚠️ 这里改成：如果你配了 expected_id，但微信回调 serial 不一致，
    # 先不直接判死刑（因为微信可能轮换公钥ID，你本地文件没更新会 mismatch）
    # 仍然尝试用现有公钥验一次；验签通过就放行。
    mismatch = bool(expected_id and serial != expected_id)

    pub_path = _get_wechatpay_public_key_path()
    if not pub_path or not os.path.exists(pub_path) or not os.path.isfile(pub_path):
        return False

    with open(pub_path, "rb") as f:
        pub_pem = f.read()
    public_key = load_pem_public_key(pub_pem)

    # message = timestamp + "\n" + nonce + "\n" + body + "\n"
    # 注意：body 必须是原始 JSON 字符串，不要做任何格式化
    try:
        body_text = raw_body.decode("utf-8")
    except Exception:
        return False

    message = f"{ts}\n{nonce}\n{body_text}\n"
    try:
        sig = base64.b64decode(signature_b64)
    except Exception:
        return False

    try:
        public_key.verify(
            sig,
            message.encode("utf-8"),
            asy_padding.PKCS1v15(),
            hashes.SHA256(),
        )
        # 如果 mismatch 但验签通过，说明你的 expected_id 可能过期/填错；
        # 这种情况下建议你把 WECHAT_PAY_PUBLIC_KEY_ID 更新成当前回调头里的 serial。
        return True
    except Exception:
        return False


def _credit_paid_order_by_notify(db, out_trade_no: str, txid: Optional[str], paid_total: Optional[int]):
    row = db.query(PayOrder).filter(PayOrder.out_trade_no == out_trade_no).first()
    if not row or row.status == "PAID":
        return

    # 金额一致性校验（建议保留）
    if paid_total is not None:
        try:
            if int(row.total_fee or 0) != int(paid_total):
                return
        except Exception:
            return

    u = db.query(User).filter(User.id == row.user_id).first()
    if not u:
        return

    u.credits = int(u.credits or 0) + int(row.credits or 0)
    row.status = "PAID"
    row.wx_transaction_id = txid or row.wx_transaction_id
    row.paid_at = datetime.utcnow()

    db.add(u)
    db.add(row)
    db.commit()


@app.post("/v1/pay/notify")
async def wechat_pay_notify(request: Request):
    """
    微信支付 v3 回调（公钥模式验签）
    - 验签失败：返回 401（微信会重试；伪造请求会被挡）
    - 验签通过：解密 resource -> 幂等入账 -> 返回 SUCCESS
    """
    raw = await request.body()

    # 1) 验签（必须）
    ok_sig = _verify_wechatpay_signature_public_key_mode(
        raw_body=raw,
        headers=dict(request.headers),
    )
    if not ok_sig:
        return JSONResponse({"code": "FAIL", "message": "SIGNATURE_ERROR"}, status_code=401)

    # 2) 解析 body
    try:
        payload = json.loads(raw.decode("utf-8") or "{}")
    except Exception:
        # 验签都通过了，body 却坏了：返回 SUCCESS 避免重试风暴，轮询兜底
        return JSONResponse({"code": "SUCCESS", "message": "OK"})

    # 3) 解密 resource（必须有 APIv3 key）
    api_v3_key = _get_api_v3_key()
    if not api_v3_key:
        # 没有 v3 key 就无法解密，但验签已通过：返回 SUCCESS，继续靠轮询兜底
        return JSONResponse({"code": "SUCCESS", "message": "OK"})

    resource = payload.get("resource") or {}
    nonce = resource.get("nonce")
    ciphertext = resource.get("ciphertext")
    associated_data = resource.get("associated_data") or ""

    if not (nonce and ciphertext):
        return JSONResponse({"code": "SUCCESS", "message": "OK"})

    try:
        plain = _aes256gcm_decrypt(api_v3_key, nonce, ciphertext, associated_data)
        data = json.loads(plain or "{}")
    except Exception:
        return JSONResponse({"code": "SUCCESS", "message": "OK"})

    # 4) 基础校验 appid/mchid（防止串号）
    env = _require_wechat_pay_env()
    if (data.get("appid") or "") != env["WECHAT_APPID"] or (data.get("mchid") or "") != env["WECHAT_MCHID"]:
        return JSONResponse({"code": "SUCCESS", "message": "OK"})

    out_trade_no = (data.get("out_trade_no") or "").strip()
    trade_state = (data.get("trade_state") or "").strip()
    txid = (data.get("transaction_id") or "").strip() or None

    paid_total = None
    try:
        paid_total = int(((data.get("amount") or {}).get("total")) or 0)
    except Exception:
        paid_total = None

    # 5) 成功单入账（幂等）
    if out_trade_no and trade_state == "SUCCESS":
        db = SessionLocal()
        try:
            _credit_paid_order_by_notify(db, out_trade_no, txid, paid_total)
        finally:
            db.close()

    return JSONResponse({"code": "SUCCESS", "message": "OK"})

@app.get("/health")
def health():
    return {"ok": True}